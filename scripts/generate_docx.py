import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_formatted_text(paragraph, html_like_text):
    # Parse bold and italic tags: <b>...</b> and <i>...</i>
    tokens = re.split(r'(<b>.*?</b>|<i>.*?</i>)', html_like_text)
    for token in tokens:
        if token.startswith('<b>') and token.endswith('</b>'):
            run = paragraph.add_run(token[3:-4])
            run.bold = True
        elif token.startswith('<i>') and token.endswith('</i>'):
            run = paragraph.add_run(token[3:-4])
            run.italic = True
        else:
            paragraph.add_run(token)

def generate_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Pre-process bold and italic formatting using HTML-like tags
    content = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', content)
    content = re.sub(r'\*(.*?)\*', r'<i>\1</i>', content)

    lines = content.split('\n')

    doc = Document()
    
    # Configure executive page margins (0.75 inches)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Base typography styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate body text

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Handle Code Blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # Compile code block paragraph
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.right_indent = Inches(0.2)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(12)
                
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Handle Tables
        if line.strip().startswith('|'):
            if '---' in line:
                i += 1
                continue
            in_table = True
            cols = [col.strip() for col in line.split('|')[1:-1]]
            table_rows.append(cols)
            i += 1
            continue
        elif in_table:
            # Table ended, compile it
            if table_rows:
                num_rows = len(table_rows)
                num_cols = len(table_rows[0])
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                for row_idx, row_data in enumerate(table_rows):
                    row = table.rows[row_idx]
                    for col_idx, cell_data in enumerate(row_data):
                        cell = row.cells[col_idx]
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(4)
                        p.paragraph_format.space_after = Pt(4)
                        
                        if row_idx == 0:
                            # Header cell formatting
                            set_cell_background(cell, "1E293B") # Dark slate
                            run = p.add_run(cell_data)
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White
                            run.font.size = Pt(9.5)
                        else:
                            # Body cell formatting
                            if row_idx % 2 == 1:
                                set_cell_background(cell, "F8FAFC") # Very light gray for zebra rows
                            else:
                                set_cell_background(cell, "FFFFFF")
                            add_formatted_text(p, cell_data)
                            p.runs[0].font.size = Pt(9)
                            
                # Add spacing after table
                p_after = doc.add_paragraph()
                p_after.paragraph_format.space_before = Pt(0)
                p_after.paragraph_format.space_after = Pt(8)
                
            in_table = False
            table_rows = []

        # Handle Horizontal Rules
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run("______________________________________________________________________")
            run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
            i += 1
            continue

        # Handle Headings
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(18)
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(22)
            run.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A) # Deep navy
            i += 1
            continue
        elif line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            i += 1
            continue
        elif line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(11.5)
            run.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            i += 1
            continue

        # Handle Bullet Points
        if line.strip().startswith('*') or line.strip().startswith('-'):
            bullet_text = line.strip()[1:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            add_formatted_text(p, bullet_text)
            i += 1
            continue

        # Handle Numbered Lists
        match_num = re.match(r'^\d+\.\s+(.*)', line.strip())
        if match_num:
            num_text = match_num.group(1)
            prefix = line.strip().split('.')[0]
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(4)
            add_formatted_text(p, num_text)
            i += 1
            continue

        # Handle Normal Paragraphs
        if line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            add_formatted_text(p, line.strip())

        i += 1

    doc.save(docx_path)
    print(f"Successfully generated executive Word document at: {docx_path}")

if __name__ == "__main__":
    generate_docx(
        "/Users/vasu/Documents/GitHub/gemma4-ira-companion/final_report.md",
        "/Users/vasu/Documents/GitHub/gemma4-ira-companion/final_report.docx"
    )
